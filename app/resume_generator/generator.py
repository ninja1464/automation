from docx import Document
import json
from datetime import datetime

def generate_resume(job_text):
    with open("master_profile.json") as f:
        profile = json.load(f)

    doc = Document()
    doc.add_heading(profile["personal_info"]["name"], level=1)
    doc.add_paragraph(profile["personal_info"]["email"])

    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(profile["skills"]))

    doc.add_heading("Experience", level=2)

    for exp in profile["experience"]:
        doc.add_heading(f"{exp['title']} - {exp['company']}", level=3)
        for bullet in exp["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet")

    file_name = f"generated_resume_{datetime.now().timestamp()}.docx"
    doc.save(file_name)

    return file_name